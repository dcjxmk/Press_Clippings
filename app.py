from flask import Flask, request, jsonify, render_template, send_file, redirect
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
from apscheduler.schedulers.background import BackgroundScheduler
import json
from selenium import webdriver
from selenium.webdriver.firefox.options import Options
from selenium.webdriver.firefox.service import Service
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.edge.service import Service as EdgeService
from webdriver_manager.firefox import GeckoDriverManager
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.microsoft import EdgeChromiumDriverManager
import logging
from readability import Document
import newspaper
from bs4 import BeautifulSoup
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import PageTemplate
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import tempfile
from langdetect import detect
import trafilatura
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from docx.oxml import parse_xml
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import time
import logging.handlers
from functools import wraps
from selenium.webdriver.firefox.options import Options as FirefoxOptions
import platform
import os.path
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from dotenv import load_dotenv
import requests
from urllib.parse import urlparse

# Load environment variables once, before initializing the app
load_dotenv()

# Define the log_performance decorator before it's used
def log_performance(func):
    """Decorator to log function execution time"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        logger.info(f"{func.__name__} took {end_time - start_time:.2f} seconds")
        return result
    return wrapper

# Add utility functions before Flask app initialization
def extract_title(html_content):
    """Extract title from HTML content"""
    try:
        doc = Document(html_content)
        return doc.title() if doc.title() else ''
    except:
        return ''

def get_clean_source_name(domain):
    """Extract and clean source name from domain"""
    if 'pressreader.com' in domain:
        return 'PressReader'
    domain = domain.replace('www.', '')
    parts = domain.split('.')
    return parts[0].title() if parts else domain

def convert_caps_to_small_caps(text):
    """Convert ALL CAPS words to small caps while preserving normal case words"""
    if not text:
        return text
    words = text.split()
    converted_words = []
    for word in words:
        if word.isupper() and len(word) > 1:
            converted_words.append(word.capitalize())
        else:
            converted_words.append(word)
    return ' '.join(converted_words)

def extract_title_from_url(url):
    """Extract title from URL slug for News24 articles"""
    try:
        # Parse the URL
        parsed_url = urlparse(url)
        
        # Get the path
        path = parsed_url.path
        
        # Split the path by '/'
        parts = path.split('/')
        
        # Find the part that contains the article title (usually the last part before the ID)
        for part in reversed(parts):
            # Skip the ID part (usually at the end, contains numbers)
            if part and not part.isdigit() and not part.endswith('.html'):
                # Remove the ID at the end if present (format: title-20250319)
                if '-' in part and part.split('-')[-1].isdigit():
                    part = '-'.join(part.split('-')[:-1])
                
                # Convert hyphens to spaces and capitalize words
                title = ' '.join(word.capitalize() for word in part.split('-'))
                return title
    except:
        pass
    
    return None

@log_performance
def scrape_url(url):
    """Basic URL scraping focused on getting title"""
    if not url or not isinstance(url, str):
        raise ValueError("Invalid URL")

    domain = urlparse(url).netloc
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }

    try:
        # Handle News24 URLs
        if 'news24.com' in domain:
            # For News24, prioritize extracting title from URL
            url_title = extract_title_from_url(url)
            
            # If we successfully extracted a title from the URL, use it immediately
            if url_title:
                return {
                    'headline': url_title,
                    'source': 'News24',
                    'content': '',  # No need for content as per user request
                    'url': url
                }
            
            # If URL extraction failed, try other methods
            driver = get_selenium_driver()
            try:
                # Use the more comprehensive get_news24_content function
                news24_result = get_news24_content(driver, url)
                
                if news24_result.get('success', False):
                    return {
                        'headline': convert_caps_to_small_caps(news24_result['headline']),
                        'source': 'News24',
                        'content': '',  # No need for content as per user request
                        'url': url
                    }
                
                # Fallback to simpler extraction if get_news24_content fails
                driver.get(url)
                time.sleep(2)
                
                # Updated selectors for News24 articles
                for selector in [
                    'h1.article__title', 
                    'h1.article-title', 
                    'article h1',
                    '.article-view__title h1',  # New selector
                    '.article__title h1',       # New selector
                    '.article-header h1',       # New selector
                    'header h1',                # More generic selector
                    'h1'                        # Most generic selector as last resort
                ]:
                    try:
                        element = driver.find_element(By.CSS_SELECTOR, selector)
                        title = element.text
                        if title and title.strip():
                            return {
                                'headline': convert_caps_to_small_caps(title.strip()),
                                'source': 'News24',
                                'content': '',
                                'url': url
                            }
                    except:
                        continue
                
                # Try XPath as a last resort
                try:
                    title_element = driver.find_element(By.XPATH, "//h1")
                    title = title_element.text
                    if title and title.strip():
                        return {
                            'headline': convert_caps_to_small_caps(title.strip()),
                            'source': 'News24',
                            'content': '',
                            'url': url
                        }
                except:
                    pass
                
                # If all else fails, use the title extracted from URL
                if url_title:
                    return {
                        'headline': url_title,
                        'source': 'News24',
                        'content': '',
                        'url': url
                    }
                
                return {
                    'headline': '',
                    'source': 'News24',
                    'content': '',
                    'url': url
                }
            except Exception as e:
                logger.warning(f"News24 Selenium scraping failed: {str(e)}")
                
                # If Selenium fails, use the title extracted from URL
                if url_title:
                    return {
                        'headline': url_title,
                        'source': 'News24',
                        'content': '',
                        'url': url
                    }
                
                return {
                    'headline': '',
                    'source': 'News24',
                    'content': '',
                    'url': url
                }
        
        # Handle PressReader URLs
        elif 'pressreader.com' in domain:
            # Get publication name from URL
            publication = get_pressreader_source(url)
            
            # Extract info from URL if possible
            url_info = extract_info_from_pressreader_url(url)
            
            # Use Selenium for PressReader content extraction
            driver = get_selenium_driver()
            try:
                # Use the enhanced PressReader content extraction function
                pressreader_result = scrape_pressreader_content(driver, url)
                
                if pressreader_result.get('success', False):
                    return {
                        'headline': convert_caps_to_small_caps(pressreader_result['headline']),
                        'source': publication,
                        'content': pressreader_result.get('content', ''),  # Include content if available
                        'url': url
                    }
                
                # If specialized extraction fails, try a more generic approach
                driver.get(url)
                time.sleep(2)
                
                # Try to get title from page title
                try:
                    page_title = driver.title
                    if page_title:
                        # Remove publication name if present
                        if '|' in page_title:
                            title = page_title.split('|')[0].strip()
                        else:
                            title = page_title.strip()
                            
                        if title:
                            return {
                                'headline': convert_caps_to_small_caps(title),
                                'source': publication,
                                'content': '',
                                'url': url
                            }
                except:
                    pass
                
                # Try various selectors for headline
                for selector in [
                    'h1.article-title', 
                    'h1:not(.publication)', 
                    '.headline', 
                    '.article-headline',
                    '.article__title',
                    'h1',
                    '.title'
                ]:
                    try:
                        element = driver.find_element(By.CSS_SELECTOR, selector)
                        title = element.text
                        if title and title.strip():
                            return {
                                'headline': convert_caps_to_small_caps(title.strip()),
                                'source': publication,
                                'content': '',
                                'url': url
                            }
                    except:
                        continue
                
                # Return with empty headline if all extraction methods fail
                return {
                    'headline': '',
                    'source': publication,
                    'content': '',
                    'url': url
                }
                
            except Exception as e:
                logger.warning(f"PressReader Selenium scraping failed: {str(e)}")
                return {
                    'headline': '',
                    'source': publication,
                    'content': '',
                    'url': url
                }
        
        # Generic handling for other URLs
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        title = extract_title(response.text)
        
        return {
            'headline': convert_caps_to_small_caps(title),
            'source': get_clean_source_name(domain),
            'content': '',
            'url': url
        }

    except Exception as e:
        logger.warning(f"Title extraction failed: {str(e)}")
        return {
            'headline': '',
            'source': get_clean_source_name(domain),
            'content': '',
            'url': url
        }

# Single Flask initialization and configuration
app = Flask(__name__)
app.config.update(
    DEBUG=os.getenv('FLASK_DEBUG', '0') == '1',
    SECRET_KEY=os.getenv('SECRET_KEY', 'default-secret-key'),
    SQLALCHEMY_DATABASE_URI=os.getenv('DATABASE_URL', 'sqlite:///press_clippings.db'),
    SQLALCHEMY_TRACK_MODIFICATIONS=False
)

# Initialize extensions only once
db = SQLAlchemy(app)
CORS(app)

# Single logging configuration
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.handlers.RotatingFileHandler('logs/app.log', maxBytes=1024*1024, backupCount=5),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class Clipping(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    headline = db.Column(db.String(200), nullable=True, default='')  # Allow empty headline
    source = db.Column(db.String(100), nullable=False)  # Source still required
    category = db.Column(db.String(50), nullable=False)  # Category still required
    content = db.Column(db.Text, nullable=True, default='')  # Allow empty content
    url = db.Column(db.String(500))
    date = db.Column(db.DateTime, default=datetime.utcnow)
    order = db.Column(db.Integer)

    def to_dict(self):
        return {
            'id': self.id,
            'headline': self.headline or '',  # Ensure empty string if None
            'source': self.source,
            'category': self.category,
            'content': self.content or '',  # Ensure empty string if None
            'url': self.url,
            'date': self.date.strftime('%Y-%m-%d'),
            'order': self.order
        }

with app.app_context():
    db.create_all()

# Cache for webdrivers
_browser_drivers = {}
_driver_paths = {
    'chrome': None,
    'edge': None,
    'firefox': None
}

# After the _driver_paths definition, add path caching
DRIVER_CACHE_FILE = 'drivers/driver_paths.json'

def load_cached_driver_paths():
    """Load cached driver paths from file"""
    global _driver_paths
    try:
        if os.path.exists(DRIVER_CACHE_FILE):
            with open(DRIVER_CACHE_FILE, 'r') as f:
                cached_paths = json.load(f)
                _driver_paths.update({
                    k: v for k, v in cached_paths.items() 
                    if v and os.path.exists(v)
                })
    except Exception as e:
        logging.warning(f"Could not load cached driver paths: {e}")

def save_driver_paths():
    """Save current driver paths to cache file"""
    try:
        os.makedirs(os.path.dirname(DRIVER_CACHE_FILE), exist_ok=True)
        with open(DRIVER_CACHE_FILE, 'w') as f:
            json.dump(_driver_paths, f)
    except Exception as e:
        logging.warning(f"Could not save driver paths: {e}")

# Call load_cached_driver_paths at startup
load_cached_driver_paths()

def get_cached_driver_path(browser_type):
    """Get cached driver path or find existing driver"""
    global _driver_paths
    
    if _driver_paths[browser_type] and os.path.exists(_driver_paths[browser_type]):
        return _driver_paths[browser_type]
    
    drivers_dir = os.path.join(os.getcwd(), 'drivers')
    os.makedirs(drivers_dir, exist_ok=True)
    
    possible_paths = {
        'chrome': [
            os.path.join(drivers_dir, 'chromedriver.exe'),
            os.path.join(os.path.expanduser('~'), '.wdm', 'drivers', 'chromedriver', 'win32', 'chromedriver.exe')
        ],
        'edge': [
            os.path.join(drivers_dir, 'msedgedriver.exe'),
            os.path.join(os.path.expanduser('~'), '.wdm', 'drivers', 'edgedriver', 'win64', 'msedgedriver.exe')
        ],
        'firefox': [
            os.path.join(drivers_dir, 'geckodriver.exe'),
            os.path.join(os.path.expanduser('~'), '.wdm', 'drivers', 'geckodriver', 'win64', 'geckodriver.exe')
        ]
    }
    
    for path in possible_paths[browser_type]:
        if (os.path.exists(path)):
            _driver_paths[browser_type] = path
            save_driver_paths()
            return path
    
    try:
        if browser_type == 'chrome':
            path = ChromeDriverManager().install()
        elif browser_type == 'edge':
            path = EdgeChromiumDriverManager().install()
        else:
            path = GeckoDriverManager().install()
        
        if os.path.exists(path):
            dest_path = os.path.join(drivers_dir, os.path.basename(path))
            import shutil
            shutil.copy2(path, dest_path)
            _driver_paths[browser_type] = dest_path
            save_driver_paths()
            return dest_path
            
        _driver_paths[browser_type] = path
        save_driver_paths()
        return path
    except Exception as e:
        logging.error(f"Failed to install {browser_type} driver: {e}")
        return None

def get_webdriver():
    """Get or create a cached webdriver instance with minimal initialization"""
    global _browser_drivers
    
    # Clean up any stale drivers
    for browser_name, driver in list(_browser_drivers.items()):
        try:
            driver.current_url
        except:
            try:
                driver.quit()
            except:
                pass
            _browser_drivers.pop(browser_name, None)

    # Try Firefox with proper profile configuration
    try:
        options = FirefoxOptions()
        options.add_argument('-headless')
        
        # Configure Firefox profile
        profile = webdriver.FirefoxProfile()
        profile.set_preference('dom.webdriver.enabled', False)
        profile.set_preference('useAutomationExtension', False)
        profile.set_preference('dom.popup_maximum', 0)
        profile.set_preference('privacy.trackingprotection.enabled', False)
        profile.set_preference('network.cookie.cookieBehavior', 0)
        profile.update_preferences()
        
        options.profile = profile
        
        geckodriver_path = os.path.join(os.getcwd(), 'drivers', 'geckodriver.exe')
        service = Service(executable_path=geckodriver_path, log_path='logs/geckodriver.log')
        
        driver = webdriver.Firefox(service=service, options=options)
        driver.set_page_load_timeout(30)
        driver.supports_cdp = False
        _browser_drivers['firefox'] = driver
        return driver
    except Exception as firefox_error:
        logging.warning(f"Firefox initialization failed: {str(firefox_error)}")
        raise firefox_error  # Don't fall back to Chrome since we're using Firefox

def cleanup_webdriver():
    """Clean up all cached webdriver instances"""
    global _browser_drivers
    for driver in _browser_drivers.values():
        try:
            driver.quit()
        except:
            pass
    _browser_drivers = {}

def execute_browser_script(driver, script):
    """Execute browser script with CDP for Chrome/Edge or standard script for Firefox"""
    try:
        if hasattr(driver, 'supports_cdp') and driver.supports_cdp:
            try:
                driver.execute_cdp_cmd('Page.setBypassCSP', {'enabled': True})
            except:
                pass
        
        driver.execute_script(script)
    except Exception as e:
        logging.warning(f"Error executing browser script: {e}")

def extract_info_from_pressreader_url(url):
    """Extract article information from PressReader URL"""
    try:
        # Parse the URL
        parsed_url = urlparse(url)
        
        # Get the path
        path = parsed_url.path
        
        # Split the path by '/'
        parts = [p for p in path.split('/') if p]
        
        # Extract publication date if available (format: YYYYMMDD)
        date_part = None
        for part in parts:
            if part.isdigit() and len(part) == 8:
                date_part = part
                
        # Extract article ID
        article_id = None
        for part in parts:
            if part.isdigit() and len(part) > 8:  # Longer numeric IDs are likely article IDs
                article_id = part
        
        # Try to extract title from URL structure if possible
        # This is more challenging for PressReader as they don't typically include article titles in URLs
        
        return {
            'publication_date': date_part,
            'article_id': article_id
        }
    except Exception as e:
        logging.error(f"Error extracting info from PressReader URL: {str(e)}")
        return None

def get_pressreader_source(url):
    """Extract publication name from Pressreader URL with improved path parsing"""
    try:
        if 'pressreader.com' not in url:
            return url.split('/')[2]
            
        parts = [p for p in url.split('/') if p]
        
        if 'south-africa' in parts:
            idx = parts.index('south-africa')
            if idx + 1 < len(parts):
                pub_name = parts[idx + 1]
                pub_name = pub_name.replace('-', ' ')
                
                if 'star' in pub_name.lower():
                    if 'early' in pub_name.lower():
                        return 'The Star Early Edition'
                    elif 'late' in pub_name.lower():
                        return 'The Star Late Edition'
                    return 'The Star'
                
                words = pub_name.split()
                clean_words = []
                for word in words:
                    if word.lower() not in ['south', 'africa', 'early', 'late', 'edition']:
                        clean_words.append(word.capitalize())
                
                pub_name = ' '.join(clean_words)
                
                if 'early' in parts[idx + 1].lower():
                    pub_name += ' Early Edition'
                elif 'late' in parts[idx + 1].lower():
                    pub_name += ' Late Edition'
                
                return pub_name.strip()
                
        if '/textview' in url:
            parts = url.split('/textview')[0].split('/')
            for part in parts:
                if 'edition' in part.lower() or 'star' in part.lower():
                    pub_name = part.replace('-', ' ')
                    words = pub_name.split()
                    clean_words = []
                    for word in words:
                        if word.lower() not in ['south', 'africa']:
                            clean_words.append(word.capitalize())
                    return ' '.join(clean_words)
    except Exception as e:
        logging.error(f"Error extracting PressReader source: {str(e)}")
    
    return 'Unknown Publication'

def get_news24_content(driver, url):
    """Extract content from News24 with optimized anti-bot and paywall handling"""
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    import time, random
    
    try:
        # Navigate and wait for page load
        driver.get(url)
        time.sleep(2 + random.random())
        
        # First try to get the title directly from the page title
        try:
            page_title = driver.title
            if page_title and 'News24' in page_title:
                # Remove "| News24" or similar from the end
                clean_title = page_title.split('|')[0].strip()
                if clean_title:
                    return {
                        'success': True,
                        'headline': clean_title,
                        'content': ''
                    }
        except:
            pass
        
        # Try to find the article content with various selectors
        content_paths = [
            "//div[contains(@class, 'article-body')]", 
            "//div[contains(@class, 'article__content')]",
            "//div[@id='article-body']",
            "//article",
            "//main//article",
            "//main",
            "//div[contains(@class, 'article')]",
            "//body"  # Last resort - use the entire body
        ]
        
        article = None
        for xpath in content_paths:
            try:
                article = driver.find_element(By.XPATH, xpath)
                if article:
                    break
            except:
                continue
        
        if not article:
            # If we can't find the article content, try to at least get the title from the page
            try:
                return {
                    'success': True,
                    'headline': driver.title.split('|')[0].strip(),
                    'content': ''
                }
            except:
                raise Exception("Could not find article content")

        # Get headline with broader xpath
        headline = ""
        headline_paths = [
            ".//h1",
            ".//*[contains(@class, 'headline')]",
            ".//*[contains(@class, 'article-title')]",
            ".//*[contains(@class, 'title')]",
            "//h1",  # Try to find any h1 on the page
            "//title"  # Last resort - use the page title
        ]
        
        for xpath in headline_paths:
            try:
                headline_elem = article.find_element(By.XPATH, xpath)
                if headline_elem and headline_elem.text.strip():
                    headline = headline_elem.text.strip()
                    break
            except:
                continue
                
        # If we still don't have a headline, try the page title
        if not headline:
            try:
                headline = driver.title.split('|')[0].strip()
            except:
                pass
        
        # Get paragraphs with a more specific approach
        content_paths = [
            ".//p[not(ancestor::div[contains(@class, 'subscription')])]",
            ".//div[contains(@class, 'article-text')]//p",
            ".//div[contains(@class, 'article-body')]//p"
        ]
        
        paragraphs = []
        for xpath in content_paths:
            try:
                elements = article.find_elements(By.XPATH, xpath)
                paragraphs = [p.text.strip() for p in elements if p.text.strip() and len(p.text.strip()) > 50]
                if paragraphs:
                    break
            except:
                continue
        
        if paragraphs:
            # Filter out subscription messages
            unwanted = ['subscribe', 'subscription', 'premium', 'register', 'sign in']
            filtered_paragraphs = [p for p in paragraphs if not any(x in p.lower() for x in unwanted)]
            
            if filtered_paragraphs:
                return {
                    'success': True,
                    'headline': headline or filtered_paragraphs[0],
                    'content': '\n\n'.join(filtered_paragraphs[:2])
                }
        
        raise Exception("No valid content found")
            
    except Exception as e:
        logging.error(f"News24 extraction failed: {str(e)}")
        return {'success': False}

def scrape_pressreader_content(driver, url):
    """Extract content from PressReader with enhanced selectors and error handling"""
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    import time, random
    
    try:
        # Always use textview mode for better content extraction
        if '/article/' in url:
            url = url.replace('/article/', '/textview/')
        
        # Set longer timeout for PressReader which can be slow to load
        driver.set_page_load_timeout(15)
        driver.get(url)
        
        # Add a small random delay to avoid detection
        time.sleep(1 + random.random())
        
        # First try to get the title directly from the page title
        try:
            page_title = driver.title
            if page_title:
                # Remove publication name if present (format: "Article Title | Publication Name")
                if '|' in page_title:
                    clean_title = page_title.split('|')[0].strip()
                else:
                    clean_title = page_title.strip()
                
                if clean_title:
                    headline = clean_title
        except:
            headline = ""
        
        # Enhanced JavaScript extraction with multiple selectors and better error handling
        extract_script = """
            let headline = '';
            // Try multiple headline selectors
            const headlineSelectors = [
                'h1.article-title', 
                'h1:not(.publication)', 
                '.headline', 
                '.article-headline',
                '.article__title',
                'h1',
                '.title'
            ];

            for (const selector of headlineSelectors) {
                const elem = document.querySelector(selector);
                if (elem && elem.innerText.trim()) {
                    headline = elem.innerText.trim();
                    break;
                }
            }

            // If no headline found, try to get it from the page title
            if (!headline) {
                const titleElem = document.querySelector('title');
                if (titleElem) headline = titleElem.innerText.split('|')[0].trim();
            }

            // Get paragraphs with more comprehensive selectors
            const paragraphs = [];
            const contentSelectors = [
                'article p', 
                '.article-text p', 
                '.article-body p',
                '.article__content p',
                '.article-content p',
                '.article p',
                '.body p',
                'main p',
                '.content p'
            ];

            for (const selector of contentSelectors) {
                const elements = document.querySelectorAll(selector);
                if (elements && elements.length > 0) {
                    elements.forEach(p => {
                        const text = p.innerText.trim();
                        if (text.length > 30 && 
                            !text.toLowerCase().includes('cookie') && 
                            !text.toLowerCase().includes('subscribe') &&
                            !text.toLowerCase().includes('sign in') &&
                            !text.toLowerCase().includes('register')) {
                            paragraphs.push(text);
                        }
                    });
                    
                    if (paragraphs.length > 0) break;
                }
            }

            return {headline, paragraphs};
        """
        
        result = driver.execute_script(extract_script)
        
        # Use the headline from JavaScript or fallback to the one from page title
        final_headline = result['headline'] or headline
        
        # If we have paragraphs, return success
        if result['paragraphs'] and len(result['paragraphs']) > 0:
            # If no headline was found but we have paragraphs, use the first paragraph as headline
            if not final_headline and result['paragraphs'][0]:
                final_headline = result['paragraphs'][0]
                # If we used the first paragraph as headline, remove it from content
                content_paragraphs = result['paragraphs'][1:3]
            else:
                content_paragraphs = result['paragraphs'][:2]
            
            return {
                'success': True,
                'headline': final_headline,
                'content': '\n\n'.join(content_paragraphs)
            }
        
        # If JavaScript extraction failed, try direct Selenium extraction
        if not final_headline or not result['paragraphs']:
            try:
                # Try to find headline with Selenium
                headline_selectors = [
                    "h1.article-title", 
                    "h1:not(.publication)", 
                    ".headline", 
                    ".article-headline",
                    ".article__title",
                    "h1",
                    ".title"
                ]
                
                for selector in headline_selectors:
                    try:
                        element = driver.find_element(By.CSS_SELECTOR, selector)
                        if element and element.text.strip():
                            final_headline = element.text.strip()
                            break
                    except:
                        continue
                
                # Try to find paragraphs with Selenium
                content_selectors = [
                    "article p", 
                    ".article-text p", 
                    ".article-body p",
                    ".article__content p",
                    ".article-content p",
                    ".article p",
                    ".body p",
                    "main p",
                    ".content p"
                ]
                
                paragraphs = []
                for selector in content_selectors:
                    try:
                        elements = driver.find_elements(By.CSS_SELECTOR, selector)
                        if elements:
                            paragraphs = [p.text.strip() for p in elements if p.text.strip() and len(p.text.strip()) > 30]
                            if paragraphs:
                                break
                    except:
                        continue
                
                if paragraphs:
                    # Filter out unwanted content
                    unwanted = ['cookie', 'subscribe', 'premium', 'register', 'sign in']
                    filtered_paragraphs = [p for p in paragraphs if not any(x in p.lower() for x in unwanted)]
                    
                    if filtered_paragraphs:
                        # If no headline was found but we have paragraphs, use the first paragraph as headline
                        if not final_headline and filtered_paragraphs[0]:
                            final_headline = filtered_paragraphs[0]
                            # If we used the first paragraph as headline, remove it from content
                            content_paragraphs = filtered_paragraphs[1:3]
                        else:
                            content_paragraphs = filtered_paragraphs[:2]
                        
                        return {
                            'success': True,
                            'headline': final_headline,
                            'content': '\n\n'.join(content_paragraphs)
                        }
            except Exception as selenium_error:
                logging.warning(f"PressReader Selenium extraction fallback failed: {str(selenium_error)}")
        
        # If we have a headline but no content, return just the headline
        if final_headline:
            return {
                'success': True,
                'headline': final_headline,
                'content': ''
            }
            
        raise Exception("Could not extract content from PressReader")
            
    except Exception as e:
        logging.error(f"PressReader parsing error: {str(e)}")
    
    return {'success': False, 'error': 'Could not extract content'}

def set_browser_headers(driver, headers):
    """Set browser headers based on driver type"""
    try:
        if hasattr(driver, 'execute_cdp_cmd'):
            # Chrome/Edge support CDP
            driver.execute_cdp_cmd('Network.setExtraHTTPHeaders', {'headers': headers})
        else:
            # Firefox doesn't support CDP, use profile
            if isinstance(driver, webdriver.Firefox):
                profile = webdriver.FirefoxProfile()
                for key, value in headers.items():
                    profile.set_preference('general.useragent.override', value)
                driver.profile = profile
    except Exception as e:
        logging.warning(f"Could not set browser headers: {str(e)}")

_driver = None

def get_selenium_driver():
    """Get or create singleton selenium driver"""
    global _driver
    if (_driver is None):
        options = uc.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        _driver = uc.Chrome(options=options)
    return _driver

# Add cleanup for Selenium driver
def cleanup():
    global _driver
    if (_driver):
        try:
            _driver.quit()
        except:
            pass
        _driver = None

import atexit
atexit.register(cleanup)

@app.route('/api/scrape', methods=['POST', 'OPTIONS'])
def scrape():
    if request.method == 'OPTIONS':
        return '', 204
        
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        url = data.get('url', '').strip()
        if not url:
            return jsonify({'error': 'No URL provided'}), 400
            
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
            
        result = scrape_url(url)
        return jsonify(result)
        
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logging.error(f"Error in scrape endpoint: {str(e)}", exc_info=True)
        return jsonify({'error': 'Internal server error'}), 500

@log_performance
def generate_pdf(clippings):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    doc = SimpleDocTemplate(
        temp_file.name,
        pagesize=A4,
        leftMargin=2*cm,
        rightMargin=2*cm,
        topMargin=2.5*cm,
        bottomMargin=2.5*cm
    )
    
    styles = getSampleStyleSheet()
    
    current_date = datetime.now()
    day_name = current_date.strftime('%A')
    date_str = current_date.strftime('%d %B %Y')
    title_text = f"Embassy Press Clippings – {day_name}, {date_str}"
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=14,
        leading=16,
        alignment=1,
        fontName='Microsoft Sans Serif',
        textColor=colors.black,
        underline=True,
        spaceAfter=12
    )
    
    category_style = ParagraphStyle(
        'CategoryStyle',
        parent=styles['Heading1'],
        fontSize=14,
        leading=16,
        fontName='Microsoft Sans Serif',
        textColor=colors.black,
        underline=True,
        alignment=0,
        spaceAfter=12
    )
    
    headline_style = ParagraphStyle(
        'HeadlineStyle',
        parent=styles['Normal'],
        fontSize=11,
        leading=11,
        fontName='Microsoft Sans Serif',
        bold=True,
        alignment=0,
        spaceAfter=0
    )
    
    content_style = ParagraphStyle(
        'ContentStyle',
        parent=styles['Normal'],
        fontSize=11,
        leading=11,
        fontName='Microsoft Sans Serif',
        alignment=0,
        spaceAfter=0
    )
    
    source_style = ParagraphStyle(
        'SourceStyle',
        parent=styles['Normal'],
        fontSize=11,
        leading=11,
        fontName='Microsoft Sans Serif',
        textColor=colors.blue,
        spaceAfter=0,
        alignment=0
    )
    
    pdfmetrics.registerFont(TTFont('Microsoft Sans Serif', 'C:/Windows/Fonts/micross.ttf'))
    
    story = []
    
    story.append(Paragraph(f"<b><u>{title_text}</u></b>", title_style))
    story.append(Spacer(1, 20))
    
    fixed_categories = [
        "Foreign Politics",
        "Domestic Politics",
        "Economy, Energy, Climate & Agriculture",
        "Verschiedenes",
        "Cartoon"
    ]
    
    category_groups = {cat: [] for cat in fixed_categories}
    for clipping in clippings:
        user_cat = clipping.category.strip()
        matched = False
        for fixed_cat in fixed_categories:
            if user_cat.lower() == fixed_cat.lower() or user_cat.lower() in fixed_cat.lower():
                category_groups[fixed_cat].append(clipping)
                matched = True
                break
        if not matched:
            category_groups["Verschiedenes"].append(clipping)

    for category in fixed_categories:
        articles = category_groups.get(category, [])
        
        if articles or category == "Cartoon":
            story.append(Paragraph(f"<b><u>{category}</u></b>", category_style))
            story.append(Spacer(1, 6))
            
            if articles:
                for clipping in sorted(articles, key=lambda x: x.order):
                    if clipping.headline.strip():
                        story.append(Paragraph(f"<b>{clipping.headline}</b>", headline_style))
                    
                    if clipping.content.strip():
                        story.append(Paragraph(clipping.content.replace('\n\n', ' '), content_style))
                    
                    date_str = clipping.date.strftime('%d/%m/%Y')
                    clean_source = get_clean_source_name(clipping.source)
                    if clipping.url:
                        source_text = f'<link href="{clipping.url}"><font color="blue">{clean_source}</font></link>'
                    else:
                        source_text = clean_source
                    story.append(Paragraph(f"{source_text}, {date_str}", source_style))
                    
                    story.append(Spacer(1, 12))
            
            story.append(Spacer(1, 6))
    
    try:
        doc.build(story)
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        raise
    
    return temp_file.name

@log_performance
def generate_docx(clippings):
    doc = DocxDocument()
    
    style = doc.styles['Normal']
    style.font.name = 'Microsoft Sans Serif'
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    current_date = datetime.now()
    day_name = current_date.strftime('%A')
    date_str = current_date.strftime('%d %B %Y')
    title = doc.add_paragraph()
    title_run = title.add_run(f"Embassy Press Clippings – {day_name}, {date_str}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.underline = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    fixed_categories = [
        "Foreign Politics",
        "Domestic Politics",
        "Economy, Energy, Climate & Agriculture",
        "Verschiedenes",
        "Cartoon"
    ]
    
    category_groups = {cat: [] for cat in fixed_categories}
    for clipping in clippings:
        user_cat = clipping.category.strip()
        matched = False
        for fixed_cat in fixed_categories:
            if user_cat.lower() == fixed_cat.lower() or user_cat.lower() in fixed_cat.lower():
                category_groups[fixed_cat].append(clipping)
                matched = True
                break
        if not matched:
            category_groups["Verschiedenes"].append(clipping)

    for category in fixed_categories:
        articles = category_groups.get(category, [])
        
        if articles or category == "Cartoon":
            cat_heading = doc.add_paragraph()
            cat_run = cat_heading.add_run(category)
            cat_run.bold = True
            cat_run.font.size = Pt(14)
            cat_run.font.name = 'Microsoft Sans Serif'
            cat_run.font.underline = True
            cat_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cat_heading.paragraph_format.space_after = Pt(12)
            
            if articles:
                for clipping in sorted(articles, key=lambda x: x.order):
                    if clipping.headline.strip():
                        headline = doc.add_paragraph()
                        headline_run = headline.add_run(clipping.headline)
                        headline_run.bold = True
                        headline_run.font.size = Pt(11)
                        headline_run.font.name = 'Microsoft Sans Serif'
                        headline.paragraph_format.space_after = Pt(0)
                        headline.paragraph_format.line_spacing = 1.0

                    if clipping.content.strip():
                        content = doc.add_paragraph()
                        content_run = content.add_run(clipping.content.replace('\n\n', ' '))
                        content_run.font.size = Pt(11)
                        content_run.font.name = 'Microsoft Sans Serif'
                        content.paragraph_format.space_after = Pt(0)
                        content.paragraph_format.line_spacing = 1.0
                    
                    source_para = doc.add_paragraph()
                    source_para.paragraph_format.space_after = Pt(12)
                    date_str = clipping.date.strftime('%d/%m/%Y')
                    clean_source = get_clean_source_name(clipping.source)
                    source_run = None
                    
                    if clipping.url:
                        rel_id = doc.part.relate_to(clipping.url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
                        hyperlink.set(docx.oxml.shared.qn('r:id'), rel_id)
                        
                        run = docx.oxml.shared.OxmlElement('w:r')
                        run_props = docx.oxml.shared.OxmlElement('w:rPr')
                        
                        color = docx.oxml.shared.OxmlElement('w:color')
                        color.set(docx.oxml.shared.qn('w:val'), '0000FF')
                        run_props.append(color)
                        
                        underline = docx.oxml.shared.OxmlElement('w:u')
                        underline.set(docx.oxml.shared.qn('w:val'), 'single')
                        run_props.append(underline)
                        
                        run.append(run_props)
                        text = docx.oxml.shared.OxmlElement('w:t')
                        text.text = clean_source
                        run.append(text)
                        hyperlink.append(run)
                        source_para._p.append(hyperlink)
                    else:
                        source_run = source_para.add_run(clean_source)
                    
                    source_para.add_run(f", {date_str}")
                    source_para.paragraph_format.line_spacing = 1.0
                    source_para.paragraph_format.space_after = Pt(6)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        doc.save(temp_file.name)
    except Exception as e:
        logging.error(f"Error generating Word document: {str(e)}")
        raise
    
    return temp_file.name

def purge_old_clippings():
    with app.app_context():
        cutoff_date = datetime.utcnow() - timedelta(hours=24)
        Clipping.query.filter(Clipping.date < cutoff_date).delete()
        db.session.commit()

scheduler = BackgroundScheduler()
scheduler.add_job(func=purge_old_clippings, trigger="interval", hours=1)
scheduler.start()

@app.route('/')
def index():
    return render_template('index.html')

@app.errorhandler(404)
def not_found_error(error):
    if request.path.startswith('/api/'):
        return jsonify({'error': 'Not Found'}), 404
    return render_template('index.html'), 404

@app.errorhandler(Exception)
def handle_error(error):
    logger.error(f"Unhandled error: {error}", exc_info=True)
    return jsonify({'error': str(error)}), 500

@app.route('/api/clippings', methods=['GET', 'POST'])
def handle_clippings():
    if request.method == 'POST':
        data = request.json
        headline = data['headline']
        if not data.get('isEdit', False):
            headline = convert_caps_to_small_caps(headline)
        clipping = Clipping(
            headline=headline,
            source=data['source'],
            category=data['category'],
            content=data['content'],
            url=data.get('url'),
            order=len(Clipping.query.all())
        )
        db.session.add(clipping)
        db.session.commit()
        return jsonify(clipping.to_dict())
    
    clippings = Clipping.query.order_by(Clipping.order).all()
    return jsonify([c.to_dict() for c in clippings])

@app.route('/api/clippings/reorder', methods=['POST'])
def reorder_clippings():
    new_order = request.json
    for item in new_order:
        clipping = Clipping.query.get(item['id'])
        if clipping:
            clipping.order = item['order']
    db.session.commit()
    return '', 204

@app.route('/api/clippings/<int:clipping_id>', methods=['PUT', 'DELETE'])
def handle_clipping(clipping_id):
    clipping = Clipping.query.get_or_404(clipping_id)
    
    if request.method == 'DELETE':
        db.session.delete(clipping)
        db.session.commit()
        return '', 204
    elif request.method == 'PUT':
        data = request.json
        for key, value in data.items():
            setattr(clipping, key, value)
        db.session.commit()
        return jsonify(clipping.to_dict())

@app.route('/api/export/pdf')
def export_pdf():
    clippings = Clipping.query.order_by(Clipping.order).all()
    pdf_path = generate_pdf(clippings)
    
    return send_file(
        pdf_path,
        mimetype='application/pdf',
        as_attachment=True,
        download_name='press_clippings.pdf'
    )

@app.route('/api/export/docx')
def export_docx():
    clippings = Clipping.query.order_by(Clipping.order).all()
    docx_path = generate_docx(clippings)
    
    return send_file(
        docx_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='press_clippings.docx'
    )

@app.route('/api/clippings/delete-all', methods=['DELETE'])
def delete_all_clippings():
    try:
        Clipping.query.delete()
        db.session.commit()
        return '', 204
    except Exception as e:
        logging.error(f"Error deleting all clippings: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/health')
def health_check():
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat()
    })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True)
